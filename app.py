
import streamlit as st
import openai
import pandas as pd
import fitz  # PyMuPDF
from docx import Document
import os

st.set_page_config(page_title="Agente Jurídico Trabalhista", layout="wide")
st.title("🤖 Agente Jurídico - Defesa Trabalhista Empresarial")

openai.api_key = st.secrets["OPENAI_API_KEY"]

# Upload dos arquivos
uploaded_pdf = st.file_uploader("📄 Envie a petição inicial (PDF)", type="pdf")
uploaded_reuniao = st.file_uploader("🗣️ Resumo da reunião (opcional)", type=["txt", "pdf"])

# Funções auxiliares
def extract_text_from_pdf(file):
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

def carregar_base(nome_arquivo):
    caminho = os.path.join("base", nome_arquivo)
    if nome_arquivo.endswith(".xlsx"):
        return pd.read_excel(caminho)
    elif nome_arquivo.endswith(".docx"):
        doc = Document(caminho)
        return "\n".join([p.text for p in doc.paragraphs])
    elif nome_arquivo.endswith(".pdf"):
        with fitz.open(caminho) as pdf:
            return "".join([page.get_text() for page in pdf])
    else:
        return ""

def montar_prompt(texto_peticao, resumo, jurisprudencia, perguntas, enunciados, cct, teses):
    return f"""Você é um especialista jurídico em defesa trabalhista empresarial.

📄 PETIÇÃO INICIAL:
{texto_peticao}

🗣️ RESUMO DA REUNIÃO:
{resumo}

📚 JURISPRUDÊNCIAS:
{jurisprudencia}

📌 ENUNCIADOS TRT7:
{enunciados}

📄 CONVENÇÃO COLETIVA APLICÁVEL:
{cct}

🤔 PERGUNTAS ESTRATÉGICAS:
{perguntas}

📘 TESES DE DEFESA PADRÃO:
{teses}

Gere um RELATÓRIO COMPLETO abordando:
- Nome da parte
- Empresa
- Advogado(a)
- Número do processo
- Detalhamento dos fatos
- Espelho dos pedidos
- Riscos
- Aplicação da CCT
- Jurisprudência
- Perguntas ao cliente
- Teses sugeridas
- Checklist de documentos necessários
"""

# Execução principal
if uploaded_pdf is not None:
    with st.spinner("🔍 Analisando documentos..."):
        texto_peticao = extract_text_from_pdf(uploaded_pdf)
        resumo_reuniao = extract_text_from_pdf(uploaded_reuniao) if uploaded_reuniao else "Não informado"
        jurisprudencia = carregar_base("Julgados - TST.pdf")
        enunciados = carregar_base("Enunciados trt 7.pdf")
        perguntas = carregar_base("Perguntas trabalhistas processo judicial.docx")
        cct = carregar_base("CCT varejo itapipoca 2025.pdf")
        teses = carregar_base("planilha de teses e perguntas para contestação - cópia.xlsx")

        prompt = montar_prompt(texto_peticao, resumo_reuniao, jurisprudencia, perguntas, enunciados, cct, teses)

        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=3500
        )
        relatorio = response.choices[0].message.content

        st.download_button("📥 Baixar relatório .txt", relatorio, file_name="relatorio_trabalhista.txt")
        docx_path = "/mnt/data/relatorio_trabalhista.docx"
        doc = Document()
        for linha in relatorio.split("\n"):
            doc.add_paragraph(linha)
        doc.save(docx_path)
        with open(docx_path, "rb") as file:
            st.download_button("📥 Baixar relatório .docx", file, file_name="relatorio_trabalhista.docx")
        st.text_area("📋 Relatório:", relatorio, height=600)
